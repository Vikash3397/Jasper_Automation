#!/usr/bin/env python3
"""Convert functional_spec/*.md spec to a user-friendly Word document (.docx)."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_IN = ROOT / "functional_spec" / "Invoice_Functional_Template.md"
DEFAULT_OUT = ROOT / "functional_spec" / "Invoice_Functional_Template.docx"


def _strip_md_inline(text: str) -> str:
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text.strip()


def _add_formatted_run(paragraph, text: str) -> None:
    """Add text with **bold** and `code` spans."""
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


def _set_cell_shading(cell, fill: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        from docx.oxml import OxmlElement

        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _add_table(doc: Document, rows: list[list[str]], header: bool = True) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell_text = _strip_md_inline(row[j]) if j < len(row) else ""
            cell = table.rows[i].cells[j]
            cell.text = cell_text
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    r.font.size = Pt(10)
            if header and i == 0:
                _set_cell_shading(cell, "D9E2F3")
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
    doc.add_paragraph()


def _parse_table_block(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        if not line.strip().startswith("|"):
            continue
        if re.match(r"^\|[\s\-:|]+\|$", line.strip()):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def convert(md_path: Path, out_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Title block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("National Invoice Usage Template")
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Functional & Technical Specification")
    sr.font.size = Pt(14)
    sr.italic = True

    doc.add_paragraph()

    how = doc.add_paragraph()
    how.add_run("How to read this document").bold = True
    doc.add_paragraph(
        "Part 1 — Functional specification: what appears on the invoice (labels, sections "
        "H1–H3, D1, F1–F2, and business rules). Part 2 — Technical specification: how the "
        "Jasper template implements Part 1 (JRXML files, queries, parameters, bands, groups, "
        "and validation). Business readers use Part 1 only; developers use Part 1 and Part 2.",
        style="Intense Quote",
    )

    i = 0
    in_code = False
    code_lines: list[str] = []
    table_buf: list[str] = []
    list_buf: list[str] = []

    def flush_table() -> None:
        nonlocal table_buf
        if table_buf:
            _add_table(doc, _parse_table_block(table_buf))
            table_buf = []

    def flush_list() -> None:
        nonlocal list_buf
        for item in list_buf:
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_run(p, item)
        list_buf = []

    while i < len(lines):
        line = lines[i]

        if in_code:
            if line.strip().startswith("```"):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(8)
                in_code = False
                code_lines = []
            else:
                code_lines.append(line)
            i += 1
            continue

        if line.strip().startswith("```"):
            flush_table()
            flush_list()
            in_code = True
            code_lines = []
            i += 1
            continue

        if line.strip().startswith("|"):
            flush_list()
            table_buf.append(line)
            i += 1
            continue
        flush_table()

        if line.strip() == "---":
            flush_list()
            doc.add_paragraph()
            i += 1
            continue

        if line.startswith("# "):
            flush_list()
            h = doc.add_heading(_strip_md_inline(line[2:]), level=0)
            h.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            i += 1
            continue

        if line.startswith("## "):
            flush_list()
            doc.add_heading(_strip_md_inline(line[3:]), level=1)
            i += 1
            continue

        if line.startswith("### "):
            flush_list()
            doc.add_heading(_strip_md_inline(line[4:]), level=2)
            i += 1
            continue

        if line.strip().startswith("- "):
            list_buf.append(line.strip()[2:])
            i += 1
            continue

        if re.match(r"^\s*-\s+\[\s\]", line):
            flush_list()
            p = doc.add_paragraph(style="List Bullet")
            p.add_run("☐ " + _strip_md_inline(line.split("]", 1)[-1].strip()))
            i += 1
            continue

        if not line.strip():
            flush_list()
            i += 1
            continue

        flush_list()
        p = doc.add_paragraph()
        _add_formatted_run(p, line)
        i += 1

    flush_table()
    flush_list()

    # Footer note
    doc.add_page_break()
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fn = fp.add_run(
        f"Generated from {md_path.name}. "
        "Machine-readable source: same folder, .md file. "
        "Spec version 1.0."
    )
    fn.font.size = Pt(9)
    fn.italic = True
    fn.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Wrote {out_path}")


def main() -> int:
    md = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_IN
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else DEFAULT_OUT
    if not md.exists():
        print(f"Missing: {md}", file=sys.stderr)
        return 1
    convert(md, out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
